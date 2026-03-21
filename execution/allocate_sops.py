#!/usr/bin/env python3
"""
Script: allocate_sops.py
Purpose: Auto-classify business SOPs/training files and route each one to the correct agent's skills.md
Inputs:  Files in uploads/ OR any directory via --source-dir (.md, .txt, .pdf, .docx)
Outputs: Each file's content reference appended to the correct agent's skills.md section
         + receipt printed to stdout showing what was allocated where
"""

import argparse
import os
import sys
import json
import shutil
from pathlib import Path
from datetime import datetime

try:
    import anthropic
except ImportError:
    print("[allocate] ERROR: anthropic package not installed. Run: pip install anthropic")
    sys.exit(1)

# ─── Paths ────────────────────────────────────────────────────────────────────

PROJECT_ROOT  = Path("/Users/sabbojb/.gemini/antigravity/playground/nomad-nebula")
MEMORY_ROOT   = Path("/Users/Shared/antigravity/memory")
UPLOADS_DIR   = MEMORY_ROOT / "uploads"
PROCESSED_DIR = UPLOADS_DIR / "processed"

AGENT_SKILLS_MAP = {
    "ads":         PROJECT_ROOT / "bots" / "ads-copy" / "skills.md",
    "copywriting": PROJECT_ROOT / "bots" / "ads-copy" / "skills.md",
    "vsl":         PROJECT_ROOT / "bots" / "content"  / "skills.md",
    "funnel":      PROJECT_ROOT / "bots" / "content"  / "skills.md",
    "content":     PROJECT_ROOT / "bots" / "content"  / "skills.md",
    "lead-gen":    PROJECT_ROOT / "bots" / "ads-copy" / "skills.md",
    "outreach":    PROJECT_ROOT / "bots" / "outreach" / "skills.md",
    "sales":       PROJECT_ROOT / "bots" / "outreach" / "skills.md",
    "closing":     PROJECT_ROOT / "bots" / "outreach" / "skills.md",
    "amazon":      MEMORY_ROOT  / "amazon" / "references.md",
    "fba":         MEMORY_ROOT  / "amazon" / "references.md",
    "fulfillment": MEMORY_ROOT  / "amazon" / "references.md",
    "coaching":    MEMORY_ROOT  / "amazon" / "references.md",
    "general":     MEMORY_ROOT  / "global" / "references.md",
}

AGENT_NAMES = {
    "ads":         "ads-copy-agent",
    "copywriting": "ads-copy-agent",
    "vsl":         "content-agent",
    "funnel":      "content-agent",
    "content":     "content-agent",
    "lead-gen":    "lead-gen-agent",
    "outreach":    "outreach-agent",
    "sales":       "outreach-agent",
    "closing":     "outreach-agent",
    "amazon":      "amazon-agent",
    "fba":         "amazon-agent",
    "fulfillment": "amazon-agent",
    "coaching":    "amazon-agent",
    "general":     "ceo (global context)",
}

VALID_DOMAINS  = list(AGENT_SKILLS_MAP.keys())
SUPPORTED_EXTS = {".md", ".txt", ".pdf", ".docx"}


# ─── File Reading ─────────────────────────────────────────────────────────────

def read_file(path: Path) -> str:
    """Read .md, .txt, .pdf, or .docx. Returns empty string on failure."""
    suffix = path.suffix.lower()

    if suffix in (".md", ".txt"):
        return path.read_text(encoding="utf-8", errors="ignore")

    elif suffix == ".pdf":
        try:
            import pypdf
            reader = pypdf.PdfReader(str(path))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
            return text
        except ImportError:
            print(f"[allocate] WARNING: pypdf not installed — skipping {path.name}")
            print("[allocate]   Install with: pip install pypdf")
            return ""
        except Exception as e:
            print(f"[allocate] WARNING: Could not read PDF {path.name}: {e}")
            return ""

    elif suffix == ".docx":
        try:
            from docx import Document
            doc = Document(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also grab text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            return "\n".join(paragraphs)
        except ImportError:
            print(f"[allocate] WARNING: python-docx not installed — skipping {path.name}")
            print("[allocate]   Install with: pip install python-docx")
            return ""
        except Exception as e:
            print(f"[allocate] WARNING: Could not read DOCX {path.name}: {e}")
            return ""

    else:
        print(f"[allocate] Unsupported type: {path.suffix} — skipping {path.name}")
        return ""


# ─── Classification ───────────────────────────────────────────────────────────

def classify_file(filename: str, content: str) -> tuple[list[str], str]:
    """Use Claude Haiku to classify into domains. Returns (domains, summary)."""
    client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY", ""))
    preview = content[:4000]

    prompt = f"""You are classifying a business training document to route it to the correct AI agent.

Filename: {filename}

Document preview:
{preview}

Assign 1–3 domain tags from this list ONLY:
- ads (Meta/YouTube/Google ads, paid advertising, ad creative, hooks)
- copywriting (general copywriting, persuasion, direct response, frameworks)
- vsl (video sales letters, VSL scripts, sales video structure)
- funnel (landing pages, funnels, conversion optimization, opt-in pages)
- content (organic content, social media, content strategy, YouTube, TikTok, Instagram)
- lead-gen (lead generation, prospecting, scraping, databases)
- outreach (cold email, DMs, Dream 100, personalized outreach, setting scripts)
- sales (sales calls, objection handling, proposals, setter/closer frameworks)
- closing (closing techniques, trial closes, tie-downs, contract negotiation)
- amazon (Amazon FBA, product research, PPC, listing optimization)
- fba (FBA logistics, fulfillment, inventory, suppliers)
- fulfillment (program delivery, student success, coaching delivery)
- coaching (coaching methodology, client results, curriculum)
- general (business strategy, mindset, general marketing — applies to all agents)

Rules:
- Pick the 1–3 most relevant tags. Do not over-tag.
- If it's clearly multi-domain include both (e.g. a closing script that's also outreach = ["closing","outreach"]).
- If ambiguous, use "general".

Respond with valid JSON only:
{{
  "domains": ["tag1", "tag2"],
  "summary": "One sentence describing what this document covers and its primary use case"
}}"""

    try:
        message = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=256,
            messages=[{"role": "user", "content": prompt}]
        )
        raw = message.content[0].text.strip()
        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"):
                raw = raw[4:]
        result = json.loads(raw.strip())
        domains = [d for d in result.get("domains", []) if d in VALID_DOMAINS]
        if not domains:
            domains = ["general"]
        summary = result.get("summary", "No summary available.")
        return domains, summary
    except Exception as e:
        print(f"[allocate] WARNING: Classification failed for {filename}: {e}")
        return ["general"], "Classification failed — routed to global memory."


# ─── Skills.md Update ─────────────────────────────────────────────────────────

def already_allocated(skills_path: Path, filename: str) -> bool:
    """Return True if this filename already has an entry in the skills file."""
    if not skills_path.exists():
        return False
    text = skills_path.read_text(encoding="utf-8", errors="ignore")
    return f"## [{filename}]" in text


def append_to_skills(skills_path: Path, filename: str, file_path: Path,
                     summary: str, domain: str, source_dir: bool) -> bool:
    """Append a reference entry to the target skills.md or references.md."""
    if already_allocated(skills_path, filename):
        print(f"[allocate]   Already in {skills_path.name} — skipping duplicate")
        return False

    if not skills_path.exists():
        skills_path.parent.mkdir(parents=True, exist_ok=True)
        skills_path.write_text("# Skills / References\n\n", encoding="utf-8")

    date = datetime.now().strftime("%Y-%m-%d")
    # For source-dir files, link back to original path; for uploads, link to processed
    if source_dir:
        location = str(file_path)
    else:
        location = f"uploads/processed/{filename}"

    entry = f"""
---

## [{filename}] — Allocated {date}
**Domain:** {domain}
**Summary:** {summary}
**Source:** `{location}`
**Usage:** Reference this document when handling [{domain}] tasks. Read the source file for full content.
"""
    with open(skills_path, "a", encoding="utf-8") as f:
        f.write(entry)
    return True


# ─── Process a single file ────────────────────────────────────────────────────

def process_file(path: Path, source_dir_mode: bool = False, dry_run: bool = False) -> dict:
    """Classify and route a single file to agent skills."""
    print(f"[allocate] Processing: {path.name}")
    content = read_file(path)
    if not content.strip():
        print(f"[allocate]   Empty/unreadable — skipping")
        return {"file": path.name, "status": "skipped", "reason": "empty or unreadable"}

    domains, summary = classify_file(path.name, content)
    routed_to = []

    if dry_run:
        for domain in domains:
            agent = AGENT_NAMES.get(domain, "global")
            routed_to.append(f"{agent} ({domain})")
        print(f"[allocate]   DRY RUN — would route to: {', '.join(routed_to)}")
        print(f"[allocate]   Summary: {summary}")
        return {"file": path.name, "status": "dry-run", "domains": domains,
                "summary": summary, "routed_to": routed_to}

    # Deduplicate target files
    seen_paths = set()
    for domain in domains:
        target = AGENT_SKILLS_MAP.get(domain, MEMORY_ROOT / "global" / "references.md")
        target_key = str(target)
        if target_key not in seen_paths:
            seen_paths.add(target_key)
            append_to_skills(target, path.name, path, summary, domain, source_dir_mode)
            agent = AGENT_NAMES.get(domain, "global")
            routed_to.append(f"{agent} ({domain})")

    # Archive only if coming from uploads/ (not --source-dir)
    if not source_dir_mode:
        PROCESSED_DIR.mkdir(exist_ok=True)
        dest = PROCESSED_DIR / path.name
        if dest.exists():
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            dest = PROCESSED_DIR / f"{path.stem}_{ts}{path.suffix}"
        shutil.move(str(path), str(dest))
        location_note = f"Archived → uploads/processed/{dest.name}"
    else:
        location_note = f"Original kept at {path}"

    print(f"[allocate] ✓ {path.name}")
    print(f"[allocate]   Domains:  {', '.join(domains)}")
    print(f"[allocate]   Summary:  {summary}")
    print(f"[allocate]   Routed →  {', '.join(routed_to)}")
    print(f"[allocate]   {location_note}")

    return {
        "file": path.name,
        "status": "allocated",
        "domains": domains,
        "summary": summary,
        "routed_to": routed_to,
    }


# ─── Main ─────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Auto-classify SOPs and route them to the correct agent skills.md"
    )
    parser.add_argument("--file",       help="Process a single specific file")
    parser.add_argument("--source-dir", help="Read files from this directory (originals NOT moved)")
    parser.add_argument("--dry-run",    action="store_true", help="Classify only — don't write anything")
    args = parser.parse_args()

    # Load .env
    env_file = PROJECT_ROOT / ".env"
    if env_file.exists():
        for line in env_file.read_text().splitlines():
            line = line.strip()
            if line and not line.startswith("#") and "=" in line:
                k, v = line.split("=", 1)
                os.environ.setdefault(k.strip(), v.strip())

    if not os.environ.get("ANTHROPIC_API_KEY"):
        print("[allocate] ERROR: ANTHROPIC_API_KEY not set in .env or environment")
        sys.exit(1)

    source_dir_mode = bool(args.source_dir)

    # ── Single file mode ──
    if args.file:
        path = Path(args.file)
        if not path.exists():
            print(f"[allocate] ERROR: File not found: {path}", file=sys.stderr)
            sys.exit(1)
        process_file(path, source_dir_mode=source_dir_mode, dry_run=args.dry_run)
        return

    # ── Directory mode ──
    if args.source_dir:
        source = Path(args.source_dir)
        if not source.exists():
            print(f"[allocate] ERROR: Source directory not found: {source}", file=sys.stderr)
            sys.exit(1)
        files = [
            f for f in source.iterdir()
            if f.is_file()
            and f.suffix.lower() in SUPPORTED_EXTS
            and not f.name.startswith(".")
        ]
    else:
        UPLOADS_DIR.mkdir(parents=True, exist_ok=True)
        files = [
            f for f in UPLOADS_DIR.iterdir()
            if f.is_file()
            and f.suffix.lower() in SUPPORTED_EXTS
            and not f.name.startswith(".")
        ]

    if not files:
        src = args.source_dir or str(UPLOADS_DIR)
        print(f"[allocate] No supported files found in {src}")
        print(f"[allocate] Supported: {', '.join(sorted(SUPPORTED_EXTS))}")
        sys.exit(0)

    print(f"[allocate] Found {len(files)} file(s) to process.")
    if source_dir_mode:
        print(f"[allocate] Source-dir mode: originals will NOT be moved.\n")
    print()

    results = []
    for f in sorted(files):
        result = process_file(f, source_dir_mode=source_dir_mode, dry_run=args.dry_run)
        results.append(result)
        print()

    # ── Receipt ──
    allocated = [r for r in results if r["status"] == "allocated"]
    skipped   = [r for r in results if r["status"] == "skipped"]
    dry_ran   = [r for r in results if r["status"] == "dry-run"]

    print("─" * 60)
    print(f"ALLOCATION RECEIPT — {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    print("─" * 60)
    print(f"  Processed: {len(results)} files")
    if dry_ran:
        print(f"  Dry-run:   {len(dry_ran)} (no writes)")
    else:
        print(f"  Allocated: {len(allocated)}")
    print(f"  Skipped:   {len(skipped)}")
    print()

    agent_counts: dict[str, int] = {}
    for r in (allocated or dry_ran):
        for agent in r.get("routed_to", []):
            agent_counts[agent] = agent_counts.get(agent, 0) + 1
    if agent_counts:
        print("  Routing summary:")
        for agent, count in sorted(agent_counts.items(), key=lambda x: -x[1]):
            print(f"    {agent}: {count} file(s)")

    if skipped:
        print()
        print("  Skipped:")
        for r in skipped:
            print(f"    {r['file']} — {r.get('reason', 'unknown')}")

    print("─" * 60)
    if not args.dry_run:
        print("[allocate] Done. Agent skills.md files updated.")
    else:
        print("[allocate] Dry run complete. No files written.")


if __name__ == "__main__":
    main()
