#!/usr/bin/env python3
"""
Script: upload_growth_os_gdoc.py
Purpose: Convert the Sabbo Growth OS Brain markdown to a formatted DOCX,
         upload to Google Drive as a Google Doc with proper headings,
         chapter breaks, tables, and easy navigation.
Inputs:  bots/creators/sabbo-growth-os-brain.md + credentials.json (OAuth2)
Outputs: Google Doc in a "Sabbo Growth OS" Drive folder (link printed)
"""

import os
import re
import sys
from pathlib import Path
from io import BytesIO

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload

# ── Config ────────────────────────────────────────────────────────────────────

PROJECT_ROOT = Path(__file__).parent.parent
CREDENTIALS_FILE = PROJECT_ROOT / "credentials.json"
TOKEN_FILE = PROJECT_ROOT / "token.json"
MD_FILE = PROJECT_ROOT / "bots" / "creators" / "sabbo-growth-os-brain.md"
DOCX_OUTPUT = PROJECT_ROOT / ".tmp" / "sabbo-growth-os-brain.docx"
DRIVE_FOLDER_NAME = "Sabbo Growth OS"

SCOPES = [
    "https://www.googleapis.com/auth/drive",
    "https://www.googleapis.com/auth/documents",
]

# ── Styles ────────────────────────────────────────────────────────────────────

def setup_styles(doc):
    """Configure document styles for professional look."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Georgia'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # Heading 1 — Part headers (chapter breaks)
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Georgia'
    h1.font.size = Pt(26)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x0d, 0x47, 0xa1)  # Deep blue
    h1.paragraph_format.space_before = Pt(36)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.page_break_before = True

    # Heading 2 — Section headers
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Georgia'
    h2.font.size = Pt(18)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x1b, 0x5e, 0x20)  # Deep green
    h2.paragraph_format.space_before = Pt(24)
    h2.paragraph_format.space_after = Pt(8)

    # Heading 3 — Sub-section headers
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Georgia'
    h3.font.size = Pt(14)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    h3.paragraph_format.space_before = Pt(18)
    h3.paragraph_format.space_after = Pt(6)

    # Heading 4
    h4 = doc.styles['Heading 4']
    h4.font.name = 'Georgia'
    h4.font.size = Pt(12)
    h4.font.bold = True
    h4.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    h4.paragraph_format.space_before = Pt(12)
    h4.paragraph_format.space_after = Pt(4)

    return doc


def add_source_tag(paragraph, tag_text):
    """Add a colored [root: ...] tag inline."""
    run = paragraph.add_run(f"  {tag_text}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.italic = True


# ── Markdown Parser ───────────────────────────────────────────────────────────

def parse_markdown_to_docx(md_text, doc):
    """Parse markdown and add formatted content to the DOCX document."""
    lines = md_text.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []
    first_h1 = True

    while i < len(lines):
        line = lines[i]

        # ── Code blocks ──
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block — add as formatted paragraph
                code_text = '\n'.join(code_lines)
                if code_text.strip():
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.left_indent = Cm(1)
                    run = p.add_run(code_text)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                    # Add light gray shading
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
                    p._element.get_or_add_pPr().append(shading)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # ── Tables ──
        if line.strip().startswith('|') and '|' in line.strip()[1:]:
            if not in_table:
                in_table = True
                table_rows = []
            # Skip separator rows like |---|---|
            if re.match(r'^\s*\|[\s\-:|]+\|\s*$', line):
                i += 1
                continue
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            table_rows.append(cells)
            i += 1
            continue
        elif in_table:
            # Table ended — render it
            _render_table(doc, table_rows)
            in_table = False
            table_rows = []
            # Don't increment i — process current line normally
            continue

        stripped = line.strip()

        # ── Skip empty lines ──
        if not stripped:
            i += 1
            continue

        # ── Horizontal rules → thin line ──
        if re.match(r'^-{3,}$|^\*{3,}$|^_{3,}$', stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            # Add a bottom border as a visual separator
            pPr = p._element.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            i += 1
            continue

        # ── Headings ──
        heading_match = re.match(r'^(#{1,6})\s+(.*)', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()

            # Extract [root: ...] tags
            root_match = re.search(r'\[root:\s*([^\]]+)\]', text)
            clean_text = re.sub(r'\s*\[root:\s*[^\]]+\]', '', text).strip()
            # Remove trailing markdown bold markers
            clean_text = clean_text.strip('*').strip()

            if level == 1:
                if first_h1:
                    # Title page — no page break for the very first heading
                    p = doc.add_heading(clean_text, level=1)
                    p.paragraph_format.page_break_before = False
                    p.paragraph_format.space_before = Pt(72)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    first_h1 = False
                else:
                    p = doc.add_heading(clean_text, level=1)
            elif level == 2:
                p = doc.add_heading(clean_text, level=2)
            elif level == 3:
                p = doc.add_heading(clean_text, level=3)
            else:
                p = doc.add_heading(clean_text, level=min(level, 4))

            if root_match:
                add_source_tag(p, f"[root: {root_match.group(1)}]")

            i += 1
            continue

        # ── Blockquotes ──
        if stripped.startswith('>'):
            quote_text = stripped.lstrip('>').strip()
            # Collect multi-line blockquotes
            while i + 1 < len(lines) and lines[i + 1].strip().startswith('>'):
                i += 1
                quote_text += '\n' + lines[i].strip().lstrip('>').strip()

            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            # Add left border styling via shading
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8F5E9"/>')
            p._element.get_or_add_pPr().append(shading)
            _add_formatted_text(p, quote_text, italic=True)
            i += 1
            continue

        # ── Bullet points ──
        bullet_match = re.match(r'^(\s*)([-*+]|\d+\.)\s+(.*)', stripped)
        if bullet_match:
            indent_level = len(bullet_match.group(1)) // 2
            marker = bullet_match.group(2)
            text = bullet_match.group(3)
            is_numbered = bool(re.match(r'\d+\.', marker))

            if is_numbered:
                p = doc.add_paragraph(style='List Number')
            else:
                p = doc.add_paragraph(style='List Bullet')

            if indent_level > 0:
                p.paragraph_format.left_indent = Cm(1.27 * (indent_level + 1))

            _add_formatted_text(p, text)
            i += 1
            continue

        # ── Regular paragraph ──
        # Collect consecutive non-special lines as one paragraph
        para_lines = [stripped]
        while (i + 1 < len(lines) and
               lines[i + 1].strip() and
               not lines[i + 1].strip().startswith('#') and
               not lines[i + 1].strip().startswith('>') and
               not lines[i + 1].strip().startswith('```') and
               not lines[i + 1].strip().startswith('|') and
               not re.match(r'^(\s*)([-*+]|\d+\.)\s+', lines[i + 1].strip()) and
               not re.match(r'^-{3,}$|^\*{3,}$|^_{3,}$', lines[i + 1].strip())):
            i += 1
            para_lines.append(lines[i].strip())

        full_text = ' '.join(para_lines)
        p = doc.add_paragraph()
        _add_formatted_text(p, full_text)
        i += 1

    # Flush remaining table
    if in_table and table_rows:
        _render_table(doc, table_rows)

    return doc


def _add_formatted_text(paragraph, text, italic=False):
    """Add text with inline markdown formatting (bold, italic, code, links)."""
    # Pattern: **bold**, *italic*, `code`, [text](url), [root: ...]
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|\[root:\s*([^\]]+)\]|\[([^\]]+)\]\([^\)]+\))'

    parts = re.split(pattern, text)
    idx = 0
    while idx < len(parts):
        part = parts[idx]
        if part is None:
            idx += 1
            continue

        # Check if this is a full match group
        if idx + 1 < len(parts):
            # Bold: **text**
            bold_text = parts[idx + 1] if idx + 1 < len(parts) else None
            italic_text = parts[idx + 2] if idx + 2 < len(parts) else None
            code_text = parts[idx + 3] if idx + 3 < len(parts) else None
            root_text = parts[idx + 4] if idx + 4 < len(parts) else None
            link_text = parts[idx + 5] if idx + 5 < len(parts) else None

            if bold_text and re.match(r'^\*\*(.+?)\*\*$', part):
                run = paragraph.add_run(bold_text)
                run.bold = True
                if italic:
                    run.italic = True
                idx += 6
                continue
            elif italic_text and re.match(r'^\*(.+?)\*$', part):
                run = paragraph.add_run(italic_text)
                run.italic = True
                idx += 6
                continue
            elif code_text and re.match(r'^`(.+?)`$', part):
                run = paragraph.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xc6, 0x28, 0x28)
                idx += 6
                continue
            elif root_text and re.match(r'^\[root:', part):
                add_source_tag(paragraph, f"[root: {root_text}]")
                idx += 6
                continue
            elif link_text and re.match(r'^\[', part):
                run = paragraph.add_run(link_text)
                run.font.color.rgb = RGBColor(0x0d, 0x47, 0xa1)
                run.underline = True
                idx += 6
                continue

        # Plain text
        if part and part.strip():
            run = paragraph.add_run(part)
            if italic:
                run.italic = True
        idx += 1


def _render_table(doc, rows):
    """Render a markdown table as a formatted DOCX table."""
    if not rows:
        return

    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Light Grid Accent 1'

    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            if c_idx < num_cols:
                cell = table.cell(r_idx, c_idx)
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(cell_text.strip())
                run.font.size = Pt(10)
                run.font.name = 'Georgia'
                if r_idx == 0:
                    run.bold = True
                    # Header row shading
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E3F2FD"/>')
                    cell._element.get_or_add_tcPr().append(shading)

    # Add spacing after table
    doc.add_paragraph()


# ── Google Drive Upload ───────────────────────────────────────────────────────

def get_drive_service():
    """Authenticate with Google Drive using OAuth2 (personal account)."""
    creds = None

    # Load existing token
    if TOKEN_FILE.exists():
        creds = Credentials.from_authorized_user_file(str(TOKEN_FILE), SCOPES)

    # If no valid creds, run OAuth flow
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            if not CREDENTIALS_FILE.exists():
                print(f"\nERROR: credentials.json not found at: {CREDENTIALS_FILE}")
                sys.exit(1)
            flow = InstalledAppFlow.from_client_secrets_file(str(CREDENTIALS_FILE), SCOPES)
            creds = flow.run_local_server(port=0)

        # Save token for next run
        TOKEN_FILE.write_text(creds.to_json())
        print("  Token saved for future runs")

    return build("drive", "v3", credentials=creds)


def get_or_create_folder(drive, name):
    """Find or create a Drive folder."""
    q = f"name='{name}' and mimeType='application/vnd.google-apps.folder' and trashed=false"
    existing = drive.files().list(q=q, fields="files(id,name)").execute().get("files", [])
    if existing:
        return existing[0]["id"]
    meta = {"name": name, "mimeType": "application/vnd.google-apps.folder"}
    f = drive.files().create(body=meta, fields="id").execute()
    return f["id"]


def upload_docx_as_gdoc(drive, docx_path, title, folder_id):
    """Upload DOCX to Drive, auto-convert to Google Doc."""
    media = MediaFileUpload(
        str(docx_path),
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        resumable=True,
    )
    file_metadata = {
        "name": title,
        "mimeType": "application/vnd.google-apps.document",
        "parents": [folder_id],
    }
    f = drive.files().create(
        body=file_metadata,
        media_body=media,
        fields="id,webViewLink",
    ).execute()
    return f.get("webViewLink", f"https://docs.google.com/document/d/{f['id']}/edit")


def share_folder(drive, folder_id):
    """Make folder accessible to anyone with the link."""
    drive.permissions().create(
        fileId=folder_id,
        body={"type": "anyone", "role": "writer"},
    ).execute()


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    print("\n=== Sabbo Growth OS Brain → Google Doc ===\n")

    # 1. Read markdown
    if not MD_FILE.exists():
        print(f"ERROR: Markdown file not found: {MD_FILE}")
        sys.exit(1)

    print(f"Reading: {MD_FILE.name}")
    md_text = MD_FILE.read_text(encoding='utf-8')
    print(f"  {len(md_text):,} characters, {len(md_text.splitlines()):,} lines")

    # 2. Create DOCX
    print("\nConverting markdown to formatted DOCX...")
    doc = Document()
    setup_styles(doc)

    # Set page margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    parse_markdown_to_docx(md_text, doc)

    # Save DOCX
    DOCX_OUTPUT.parent.mkdir(exist_ok=True)
    doc.save(str(DOCX_OUTPUT))
    docx_size = DOCX_OUTPUT.stat().st_size
    print(f"  Saved: {DOCX_OUTPUT} ({docx_size:,} bytes)")

    # 3. Upload to Google Drive
    print("\nAuthenticating with Google Drive...")
    drive = get_drive_service()
    print("  Authenticated")

    print(f"\nSetting up Drive folder: '{DRIVE_FOLDER_NAME}'...")
    folder_id = get_or_create_folder(drive, DRIVE_FOLDER_NAME)
    folder_url = f"https://drive.google.com/drive/folders/{folder_id}"

    share_folder(drive, folder_id)
    print(f"  Folder shared (anyone with link can edit)")

    print("\nUploading as Google Doc...")
    doc_url = upload_docx_as_gdoc(
        drive, DOCX_OUTPUT,
        "Sabbo Growth OS Brain — Complete Operating System",
        folder_id,
    )

    print(f"\n{'='*60}")
    print("DONE!")
    print(f"{'='*60}")
    print(f"Folder: {folder_url}")
    print(f"Doc:    {doc_url}")
    print(f"{'='*60}\n")

    return doc_url


if __name__ == "__main__":
    main()
