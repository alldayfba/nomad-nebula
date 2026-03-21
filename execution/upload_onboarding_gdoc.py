#!/usr/bin/env python3
"""
Script: upload_onboarding_gdoc.py
Purpose: Convert a student onboarding markdown to a formatted DOCX,
         upload to Google Drive as a Google Doc.
Usage:   python execution/upload_onboarding_gdoc.py [path_to_md]
         Defaults to .tmp/onboarding-docs/mike-walker-onboarding.md
"""

import os
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload

# ── Config ────────────────────────────────────────────────────────────────────

PROJECT_ROOT = Path(__file__).parent.parent
CREDENTIALS_FILE = PROJECT_ROOT / "credentials.json"
TOKEN_FILE = PROJECT_ROOT / "token.json"
DRIVE_FOLDER_NAME = "24/7 Profits — Student Onboarding"

SCOPES = [
    "https://www.googleapis.com/auth/drive",
    "https://www.googleapis.com/auth/documents",
]

# ── Styles ────────────────────────────────────────────────────────────────────

def setup_styles(doc):
    """Configure document styles — clean, professional, coaching-brand feel."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Georgia'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # Heading 1 — Main title / section breaks
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Georgia'
    h1.font.size = Pt(26)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x0d, 0x47, 0xa1)
    h1.paragraph_format.space_before = Pt(36)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.page_break_before = True

    # Heading 2 — Section headers (numbered sections)
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Georgia'
    h2.font.size = Pt(18)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x1b, 0x5e, 0x20)
    h2.paragraph_format.space_before = Pt(24)
    h2.paragraph_format.space_after = Pt(8)

    # Heading 3 — Sub-sections (weeks, phases)
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Georgia'
    h3.font.size = Pt(14)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    h3.paragraph_format.space_before = Pt(18)
    h3.paragraph_format.space_after = Pt(6)

    # Heading 4
    h4 = doc.styles['Heading 4']
    h4.font.name = 'Georgia'
    h4.font.size = Pt(12)
    h4.font.bold = True
    h4.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    h4.paragraph_format.space_before = Pt(12)
    h4.paragraph_format.space_after = Pt(4)

    return doc


# ── Markdown Parser ───────────────────────────────────────────────────────────

def parse_markdown_to_docx(md_text, doc):
    """Parse markdown and add formatted content to the DOCX document."""
    lines = md_text.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []
    first_h1 = True

    while i < len(lines):
        line = lines[i]

        # ── Code blocks ──
        if line.strip().startswith('```'):
            if in_code_block:
                code_text = '\n'.join(code_lines)
                if code_text.strip():
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.left_indent = Cm(1)
                    run = p.add_run(code_text)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
                    p._element.get_or_add_pPr().append(shading)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # ── Tables ──
        if line.strip().startswith('|') and '|' in line.strip()[1:]:
            if not in_table:
                in_table = True
                table_rows = []
            if re.match(r'^\s*\|[\s\-:|]+\|\s*$', line):
                i += 1
                continue
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            table_rows.append(cells)
            i += 1
            continue
        elif in_table:
            _render_table(doc, table_rows)
            in_table = False
            table_rows = []
            continue

        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # ── Horizontal rules ──
        if re.match(r'^-{3,}$|^\*{3,}$|^_{3,}$', stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            pPr = p._element.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            i += 1
            continue

        # ── Headings ──
        heading_match = re.match(r'^(#{1,6})\s+(.*)', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip().strip('*').strip()

            if level == 1:
                if first_h1:
                    p = doc.add_heading(text, level=1)
                    p.paragraph_format.page_break_before = False
                    p.paragraph_format.space_before = Pt(72)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    first_h1 = False
                else:
                    p = doc.add_heading(text, level=1)
            else:
                p = doc.add_heading(text, level=min(level, 4))

            i += 1
            continue

        # ── Blockquotes ──
        if stripped.startswith('>'):
            quote_text = stripped.lstrip('>').strip()
            while i + 1 < len(lines) and lines[i + 1].strip().startswith('>'):
                i += 1
                quote_text += '\n' + lines[i].strip().lstrip('>').strip()

            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8F5E9"/>')
            p._element.get_or_add_pPr().append(shading)
            _add_formatted_text(p, quote_text, italic=True)
            i += 1
            continue

        # ── Bullet points ──
        bullet_match = re.match(r'^(\s*)([-*+]|\d+\.)\s+(.*)', stripped)
        if bullet_match:
            indent_level = len(bullet_match.group(1)) // 2
            marker = bullet_match.group(2)
            text = bullet_match.group(3)
            is_numbered = bool(re.match(r'\d+\.', marker))

            if is_numbered:
                p = doc.add_paragraph(style='List Number')
            else:
                p = doc.add_paragraph(style='List Bullet')

            if indent_level > 0:
                p.paragraph_format.left_indent = Cm(1.27 * (indent_level + 1))

            _add_formatted_text(p, text)
            i += 1
            continue

        # ── Regular paragraph ──
        para_lines = [stripped]
        while (i + 1 < len(lines) and
               lines[i + 1].strip() and
               not lines[i + 1].strip().startswith('#') and
               not lines[i + 1].strip().startswith('>') and
               not lines[i + 1].strip().startswith('```') and
               not lines[i + 1].strip().startswith('|') and
               not re.match(r'^(\s*)([-*+]|\d+\.)\s+', lines[i + 1].strip()) and
               not re.match(r'^-{3,}$|^\*{3,}$|^_{3,}$', lines[i + 1].strip())):
            i += 1
            para_lines.append(lines[i].strip())

        full_text = ' '.join(para_lines)
        p = doc.add_paragraph()
        _add_formatted_text(p, full_text)
        i += 1

    if in_table and table_rows:
        _render_table(doc, table_rows)

    return doc


def _add_formatted_text(paragraph, text, italic=False):
    """Add text with inline markdown formatting (bold, italic, code)."""
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)'
    parts = re.split(pattern, text)
    idx = 0
    while idx < len(parts):
        part = parts[idx]
        if part is None:
            idx += 1
            continue

        # Bold
        if re.match(r'^\*\*(.+?)\*\*$', part):
            bold_text = parts[idx + 1] if idx + 1 < len(parts) else part.strip('*')
            run = paragraph.add_run(bold_text)
            run.bold = True
            if italic:
                run.italic = True
            idx += 4
            continue
        # Italic
        elif re.match(r'^\*(.+?)\*$', part) and not part.startswith('**'):
            italic_text = parts[idx + 2] if idx + 2 < len(parts) else part.strip('*')
            run = paragraph.add_run(italic_text)
            run.italic = True
            idx += 4
            continue
        # Code
        elif re.match(r'^`(.+?)`$', part):
            code_text = parts[idx + 3] if idx + 3 < len(parts) else part.strip('`')
            run = paragraph.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xc6, 0x28, 0x28)
            idx += 4
            continue

        # Plain text
        if part and part.strip():
            run = paragraph.add_run(part)
            if italic:
                run.italic = True
        idx += 1


def _render_table(doc, rows):
    """Render a markdown table as a formatted DOCX table."""
    if not rows:
        return

    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Light Grid Accent 1'

    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            if c_idx < num_cols:
                cell = table.cell(r_idx, c_idx)
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(cell_text.strip())
                run.font.size = Pt(10)
                run.font.name = 'Georgia'
                if r_idx == 0:
                    run.bold = True
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E3F2FD"/>')
                    cell._element.get_or_add_tcPr().append(shading)

    doc.add_paragraph()


# ── Google Drive Upload ───────────────────────────────────────────────────────

def get_drive_service():
    """Authenticate with Google Drive using OAuth2."""
    creds = None

    if TOKEN_FILE.exists():
        creds = Credentials.from_authorized_user_file(str(TOKEN_FILE), SCOPES)

    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            if not CREDENTIALS_FILE.exists():
                print(f"\nERROR: credentials.json not found at: {CREDENTIALS_FILE}")
                sys.exit(1)
            flow = InstalledAppFlow.from_client_secrets_file(str(CREDENTIALS_FILE), SCOPES)
            creds = flow.run_local_server(port=0)

        TOKEN_FILE.write_text(creds.to_json())
        print("  Token saved")

    return build("drive", "v3", credentials=creds)


def get_or_create_folder(drive, name):
    """Find or create a Drive folder."""
    q = f"name='{name}' and mimeType='application/vnd.google-apps.folder' and trashed=false"
    existing = drive.files().list(q=q, fields="files(id,name)").execute().get("files", [])
    if existing:
        return existing[0]["id"]
    meta = {"name": name, "mimeType": "application/vnd.google-apps.folder"}
    f = drive.files().create(body=meta, fields="id").execute()
    return f["id"]


def upload_docx_as_gdoc(drive, docx_path, title, folder_id):
    """Upload DOCX to Drive, auto-convert to Google Doc."""
    media = MediaFileUpload(
        str(docx_path),
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        resumable=True,
    )
    file_metadata = {
        "name": title,
        "mimeType": "application/vnd.google-apps.document",
        "parents": [folder_id],
    }
    f = drive.files().create(
        body=file_metadata,
        media_body=media,
        fields="id,webViewLink",
    ).execute()
    return f.get("webViewLink", f"https://docs.google.com/document/d/{f['id']}/edit")


def share_folder(drive, folder_id):
    """Make folder accessible to anyone with the link."""
    drive.permissions().create(
        fileId=folder_id,
        body={"type": "anyone", "role": "writer"},
    ).execute()


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    # Accept optional arg for markdown path
    if len(sys.argv) > 1:
        md_file = Path(sys.argv[1])
    else:
        md_file = PROJECT_ROOT / ".tmp" / "onboarding-docs" / "mike-walker-onboarding.md"

    if not md_file.exists():
        print(f"ERROR: File not found: {md_file}")
        sys.exit(1)

    # Derive title from filename
    student_name = md_file.stem.replace('-onboarding', '').replace('-', ' ').title()
    doc_title = f"24/7 Profits — {student_name} — 90-Day Game Plan"
    docx_output = PROJECT_ROOT / ".tmp" / "onboarding-docs" / f"{md_file.stem}.docx"

    print(f"\n=== Student Onboarding → Google Doc ===\n")
    print(f"Student:  {student_name}")
    print(f"Source:   {md_file.name}")

    # 1. Read markdown
    md_text = md_file.read_text(encoding='utf-8')
    print(f"  {len(md_text):,} chars, {len(md_text.splitlines()):,} lines")

    # 2. Convert to DOCX
    print("\nConverting to formatted DOCX...")
    doc = Document()
    setup_styles(doc)

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    parse_markdown_to_docx(md_text, doc)

    docx_output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_output))
    print(f"  Saved: {docx_output} ({docx_output.stat().st_size:,} bytes)")

    # 3. Upload to Google Drive
    print("\nAuthenticating with Google Drive...")
    drive = get_drive_service()
    print("  Authenticated")

    print(f"\nDrive folder: '{DRIVE_FOLDER_NAME}'...")
    folder_id = get_or_create_folder(drive, DRIVE_FOLDER_NAME)
    folder_url = f"https://drive.google.com/drive/folders/{folder_id}"

    share_folder(drive, folder_id)
    print("  Shared (anyone with link)")

    print("\nUploading as Google Doc...")
    doc_url = upload_docx_as_gdoc(drive, docx_output, doc_title, folder_id)

    print(f"\n{'='*60}")
    print("DONE!")
    print(f"{'='*60}")
    print(f"Folder: {folder_url}")
    print(f"Doc:    {doc_url}")
    print(f"{'='*60}\n")

    return doc_url


if __name__ == "__main__":
    main()
